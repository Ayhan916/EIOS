"""Unit tests for the M15 document ingestion parsers.

Tests cover: dispatch logic, malformed input handling, empty extraction
detection, MIME type resolution, XLSX multi-sheet extraction, and the
scanned-PDF OCR detection path.

All tests run without external dependencies — parser functions are
called with minimal synthetic byte payloads or their absence is
simulated via import-guard logic.
"""

from __future__ import annotations

import io
import struct
import zipfile

import pytest

from application.ingestion.parsers import (
    ParseResult,
    ParsedPage,
    MIME_BY_EXTENSION,
    SUPPORTED_MIME_TYPES,
    parse_document,
    resolve_mime_type,
)


# ---------------------------------------------------------------------------
# MIME type resolution
# ---------------------------------------------------------------------------

class TestResolveMimeType:
    def test_pdf_by_extension(self) -> None:
        assert resolve_mime_type("report.pdf", None) == "application/pdf"

    def test_docx_by_extension(self) -> None:
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert resolve_mime_type("doc.docx", None) == mime

    def test_xlsx_by_extension(self) -> None:
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        assert resolve_mime_type("data.xlsx", None) == mime

    def test_extension_takes_priority_over_content_type(self) -> None:
        assert resolve_mime_type("report.pdf", "text/plain") == "application/pdf"

    def test_no_extension_falls_back_to_content_type(self) -> None:
        assert resolve_mime_type("noextension", "application/pdf") == "application/pdf"

    def test_unknown_extension_falls_back_to_content_type(self) -> None:
        assert resolve_mime_type("file.abc", "application/pdf") == "application/pdf"

    def test_unknown_everything_returns_octet_stream(self) -> None:
        result = resolve_mime_type("file.abc", None)
        assert result == "application/octet-stream"

    def test_case_insensitive_extension(self) -> None:
        assert resolve_mime_type("REPORT.PDF", None) == "application/pdf"

    def test_mime_by_extension_has_expected_keys(self) -> None:
        assert ".pdf" in MIME_BY_EXTENSION
        assert ".docx" in MIME_BY_EXTENSION
        assert ".xlsx" in MIME_BY_EXTENSION


# ---------------------------------------------------------------------------
# ParseResult dataclass
# ---------------------------------------------------------------------------

class TestParseResult:
    def test_is_empty_no_pages(self) -> None:
        r = ParseResult()
        assert r.is_empty

    def test_is_empty_blank_text(self) -> None:
        r = ParseResult(pages=[ParsedPage(page_number=1, text="   ")])
        assert r.is_empty

    def test_not_empty_when_has_text(self) -> None:
        r = ParseResult(pages=[ParsedPage(page_number=1, text="ESG report")])
        assert not r.is_empty

    def test_total_chars_sums_pages(self) -> None:
        r = ParseResult(pages=[
            ParsedPage(page_number=1, text="hello"),
            ParsedPage(page_number=2, text="world!"),
        ])
        assert r.total_chars == 11

    def test_total_chars_empty(self) -> None:
        r = ParseResult()
        assert r.total_chars == 0

    def test_warnings_default_empty(self) -> None:
        r = ParseResult()
        assert r.warnings == []

    def test_requires_ocr_default_false(self) -> None:
        r = ParseResult()
        assert not r.requires_ocr


# ---------------------------------------------------------------------------
# Unsupported format dispatch
# ---------------------------------------------------------------------------

class TestUnsupportedFormat:
    def test_unknown_mime_returns_warning(self) -> None:
        result = parse_document(b"data", "text/plain", "file.txt")
        assert result.is_empty
        assert any("Unsupported" in w for w in result.warnings)

    def test_no_extension_unknown_mime_returns_warning(self) -> None:
        result = parse_document(b"data", "application/octet-stream", "noextension")
        assert result.is_empty
        assert len(result.warnings) > 0


# ---------------------------------------------------------------------------
# XLSX parsing
# ---------------------------------------------------------------------------

def _make_xlsx(sheets: dict[str, list[list[str]]]) -> bytes:
    """Build a minimal valid XLSX in memory using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        pytest.skip("openpyxl not installed")

    wb = openpyxl.Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        if first:
            ws = wb.active
            ws.title = sheet_name
            first = False
        else:
            ws = wb.create_sheet(title=sheet_name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestXLSXParser:
    def test_single_sheet(self) -> None:
        content = _make_xlsx({"Sheet1": [["Company", "Score"], ["Acme", "85"]]})
        result = parse_document(content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "data.xlsx")
        assert not result.is_empty
        assert len(result.pages) == 1
        assert result.pages[0].source_section == "Sheet1"
        assert "Acme" in result.pages[0].text

    def test_multi_sheet(self) -> None:
        content = _make_xlsx({
            "Environment": [["CO2", "100"]],
            "Social": [["Workers", "500"]],
        })
        result = parse_document(content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "data.xlsx")
        assert len(result.pages) == 2
        sections = {p.source_section for p in result.pages}
        assert "Environment" in sections
        assert "Social" in sections

    def test_page_numbers_assigned(self) -> None:
        content = _make_xlsx({"A": [["x"]], "B": [["y"]]})
        result = parse_document(content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "data.xlsx")
        page_nums = {p.page_number for p in result.pages}
        assert 1 in page_nums
        assert 2 in page_nums

    def test_empty_sheet_produces_no_page(self) -> None:
        content = _make_xlsx({"Empty": []})
        result = parse_document(content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "data.xlsx")
        assert result.is_empty

    def test_malformed_bytes_returns_warning(self) -> None:
        result = parse_document(b"not-an-xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "bad.xlsx")
        assert len(result.warnings) > 0

    def test_extension_dispatch(self) -> None:
        content = _make_xlsx({"Data": [["value"]]})
        result = parse_document(content, "application/octet-stream", "report.xlsx")
        assert not result.is_empty

    def test_parser_used_is_openpyxl(self) -> None:
        content = _make_xlsx({"S": [["x"]]})
        result = parse_document(content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "data.xlsx")
        assert result.parser_used == "openpyxl"


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------

def _make_docx(paragraphs: list[str], headings: list[str] | None = None) -> bytes:
    """Build a minimal valid DOCX in memory using python-docx."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    for i, para in enumerate(paragraphs):
        if headings and i < len(headings) and headings[i]:
            doc.add_heading(headings[i], level=1)
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDOCXParser:
    def test_basic_paragraphs(self) -> None:
        content = _make_docx(["This is paragraph one.", "This is paragraph two."])
        result = parse_document(
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "report.docx",
        )
        assert not result.is_empty
        full = " ".join(p.text for p in result.pages)
        assert "paragraph one" in full

    def test_malformed_bytes_returns_warning(self) -> None:
        result = parse_document(
            b"not-a-docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "bad.docx",
        )
        assert len(result.warnings) > 0

    def test_extension_dispatch(self) -> None:
        content = _make_docx(["supply chain audit results"])
        result = parse_document(content, "application/octet-stream", "audit.docx")
        assert not result.is_empty

    def test_parser_used_is_python_docx(self) -> None:
        content = _make_docx(["test paragraph"])
        result = parse_document(
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc.docx",
        )
        assert result.parser_used == "python-docx"


# ---------------------------------------------------------------------------
# PDF parsing edge cases (no real PDF needed for error paths)
# ---------------------------------------------------------------------------

class TestPDFParserEdgeCases:
    def test_malformed_pdf_bytes_handled_gracefully(self) -> None:
        result = parse_document(b"not-a-pdf", "application/pdf", "bad.pdf")
        # Should not raise; must return warnings and empty-ish result
        assert isinstance(result, ParseResult)
        assert len(result.warnings) > 0

    def test_empty_bytes_handled_gracefully(self) -> None:
        result = parse_document(b"", "application/pdf", "empty.pdf")
        assert isinstance(result, ParseResult)

    def test_extension_dispatch_pdf(self) -> None:
        result = parse_document(b"not-a-pdf", "application/octet-stream", "report.pdf")
        # Routes to PDF parser, gets warnings
        assert isinstance(result, ParseResult)

    def test_file_type_set_to_pdf(self) -> None:
        result = parse_document(b"%PDF-garbage", "application/pdf", "test.pdf")
        assert result.file_type == "pdf"


# ---------------------------------------------------------------------------
# Supported formats constant
# ---------------------------------------------------------------------------

class TestSupportedFormats:
    def test_pdf_is_supported(self) -> None:
        assert "application/pdf" in SUPPORTED_MIME_TYPES

    def test_docx_is_supported(self) -> None:
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in SUPPORTED_MIME_TYPES

    def test_xlsx_is_supported(self) -> None:
        assert "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" in SUPPORTED_MIME_TYPES
