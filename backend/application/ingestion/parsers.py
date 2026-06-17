"""
Document parsers for the EIOS ingestion pipeline (M15).

Supported formats: PDF (primary: pypdf, fallback: pdfminer.six), DOCX, XLSX.

Design principles:
- Never raise: all errors become warnings and partial results.
- Detect scanned PDFs (empty text layer) and flag for OCR rather than fail.
- Each page/sheet is tracked separately for traceability.
- Multilingual content is preserved as-is; the embedding model handles it.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

SUPPORTED_MIME_TYPES: frozenset[str] = frozenset(
    {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        # Browser variants
        "application/msword",
        "application/vnd.ms-excel",
    }
)

MIME_BY_EXTENSION: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".doc": "application/msword",
    ".xls": "application/vnd.ms-excel",
}

_WHITESPACE_RE = re.compile(r"\s+")


@dataclass
class ParsedPage:
    """Text extracted from a single logical unit of a document."""

    page_number: int  # 1-indexed; 0 for formats without page concepts
    text: str
    source_section: str | None = None  # worksheet name (XLSX) or heading/style (DOCX)


@dataclass
class ParseResult:
    pages: list[ParsedPage] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    requires_ocr: bool = False
    parser_used: str = ""
    file_type: str = ""

    @property
    def is_empty(self) -> bool:
        return not any(p.text.strip() for p in self.pages)

    @property
    def total_chars(self) -> int:
        return sum(len(p.text) for p in self.pages)


def resolve_mime_type(filename: str, content_type: str | None) -> str:
    """Return the best-guess MIME type from filename extension, falling back to content_type."""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return MIME_BY_EXTENSION.get(ext, content_type or "application/octet-stream")


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------


def _parse_pdf_pypdf(content: bytes) -> tuple[list[ParsedPage], list[str], bool]:
    """Primary PDF parser using pypdf. Returns (pages, warnings, requires_ocr)."""
    import pypdf  # noqa: PLC0415 — optional heavy dep, import on demand

    warnings: list[str] = []
    pages: list[ParsedPage] = []

    try:
        reader = pypdf.PdfReader(io.BytesIO(content), strict=False)
    except Exception as exc:
        return [], [f"pypdf could not open PDF: {exc}"], False

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            return [], ["PDF is password-protected and could not be decrypted"], False

    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
            # Normalise ligatures and Unicode whitespace
            text = _WHITESPACE_RE.sub(" ", text).strip()
        except Exception as exc:
            warnings.append(f"Page {i + 1}: pypdf extraction error — {exc}")
            text = ""
        pages.append(ParsedPage(page_number=i + 1, text=text))

    total_text = " ".join(p.text for p in pages)
    requires_ocr = bool(pages) and not total_text.strip()
    return pages, warnings, requires_ocr


def _parse_pdf_pdfminer(content: bytes) -> tuple[list[ParsedPage], list[str]]:
    """Fallback PDF parser using pdfminer.six. Handles more edge cases for older PDFs."""
    try:
        from pdfminer.high_level import extract_pages  # noqa: PLC0415
        from pdfminer.layout import LTTextContainer  # noqa: PLC0415
    except ImportError:
        return [], ["pdfminer.six not installed; fallback unavailable"]

    warnings: list[str] = []
    pages: list[ParsedPage] = []

    try:
        for i, page_layout in enumerate(extract_pages(io.BytesIO(content))):
            page_text_parts: list[str] = []
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    try:
                        page_text_parts.append(element.get_text())
                    except Exception as exc:
                        warnings.append(f"Page {i + 1}: pdfminer element error — {exc}")
            text = _WHITESPACE_RE.sub(" ", " ".join(page_text_parts)).strip()
            pages.append(ParsedPage(page_number=i + 1, text=text))
    except Exception as exc:
        return [], [f"pdfminer extraction failed: {exc}"]

    return pages, warnings


def parse_pdf(content: bytes) -> ParseResult:
    """Parse PDF with pypdf; fall back to pdfminer.six if the text layer is empty."""
    pages, warnings, requires_ocr = _parse_pdf_pypdf(content)

    if pages and requires_ocr:
        # Attempt fallback before declaring OCR required
        fallback_pages, fallback_warnings = _parse_pdf_pdfminer(content)
        fallback_text = " ".join(p.text for p in fallback_pages)
        if fallback_text.strip():
            return ParseResult(
                pages=fallback_pages,
                warnings=warnings + fallback_warnings,
                parser_used="pdfminer",
                file_type="pdf",
            )
        # Both parsers returned empty — scanned document
        return ParseResult(
            pages=pages,
            warnings=warnings
            + fallback_warnings
            + ["PDF appears to be a scanned image with no text layer. OCR processing is required."],
            requires_ocr=True,
            parser_used="pypdf+pdfminer(failed)",
            file_type="pdf",
        )

    if not pages:
        # pypdf failed entirely — try fallback
        fallback_pages, fallback_warnings = _parse_pdf_pdfminer(content)
        return ParseResult(
            pages=fallback_pages,
            warnings=warnings + fallback_warnings,
            parser_used="pdfminer",
            file_type="pdf",
        )

    return ParseResult(pages=pages, warnings=warnings, parser_used="pypdf", file_type="pdf")


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------


def parse_docx(content: bytes) -> ParseResult:
    """Parse DOCX using python-docx. Extracts paragraphs and table cells."""
    try:
        from docx import Document  # noqa: PLC0415
        from docx.opc.exceptions import PackageNotFoundError  # noqa: PLC0415
    except ImportError:
        return ParseResult(warnings=["python-docx not installed"])

    warnings: list[str] = []
    parts: list[str] = []
    current_heading: str | None = None
    section_pages: list[ParsedPage] = []

    try:
        doc = Document(io.BytesIO(content))
    except (PackageNotFoundError, Exception) as exc:
        return ParseResult(warnings=[f"DOCX could not be opened: {exc}"])

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style_name.lower().startswith("heading"):
            # Flush previous section
            if parts:
                section_text = " ".join(parts)
                section_pages.append(
                    ParsedPage(page_number=0, text=section_text, source_section=current_heading)
                )
                parts = []
            current_heading = text
        else:
            parts.append(text)

    # Flush last section
    if parts:
        section_pages.append(
            ParsedPage(page_number=0, text=" ".join(parts), source_section=current_heading)
        )

    # Extract table cells
    table_parts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                table_parts.append(" | ".join(row_texts))

    if table_parts:
        section_pages.append(
            ParsedPage(page_number=0, text="\n".join(table_parts), source_section="Tables")
        )

    if not section_pages:
        warnings.append("DOCX contained no extractable text")

    return ParseResult(
        pages=section_pages, warnings=warnings, parser_used="python-docx", file_type="docx"
    )


# ---------------------------------------------------------------------------
# XLSX parsing
# ---------------------------------------------------------------------------


def parse_xlsx(content: bytes) -> ParseResult:
    """Parse XLSX using openpyxl. Each worksheet becomes a separate page."""
    try:
        import openpyxl  # noqa: PLC0415
    except ImportError:
        return ParseResult(warnings=["openpyxl not installed"])

    warnings: list[str] = []
    pages: list[ParsedPage] = []

    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        return ParseResult(warnings=[f"XLSX could not be opened: {exc}"])

    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        try:
            ws = wb[sheet_name]
            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cell_texts = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cell_texts:
                    rows.append(" | ".join(cell_texts))
            if rows:
                pages.append(
                    ParsedPage(
                        page_number=sheet_idx + 1,
                        text="\n".join(rows),
                        source_section=sheet_name,
                    )
                )
        except Exception as exc:
            warnings.append(f"Worksheet '{sheet_name}': extraction error — {exc}")

    wb.close()

    if not pages:
        warnings.append("XLSX contained no extractable data")

    return ParseResult(pages=pages, warnings=warnings, parser_used="openpyxl", file_type="xlsx")


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------


def parse_document(content: bytes, mime_type: str, filename: str) -> ParseResult:
    """Route to the correct parser based on MIME type."""
    resolved = resolve_mime_type(filename, mime_type)

    if resolved == "application/pdf":
        return parse_pdf(content)

    if resolved == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx(content)

    if resolved == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return parse_xlsx(content)

    return ParseResult(
        warnings=[
            f"Unsupported file type '{resolved}' for file '{filename}'. Supported: PDF, DOCX, XLSX."
        ]
    )
